import sys
import os
import json
import shutil
import tempfile
from datetime import datetime
import concurrent.futures
from functools import partial
import fnmatch

from PyQt6.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QLineEdit,
                             QLabel, QFileDialog, QTextEdit, QListWidget, QMessageBox, QStyle, QStyleFactory,
                             QProgressBar, QTableWidget, QTableWidgetItem, QHeaderView, QDialogButtonBox,
                             QMainWindow, QToolBar, QAbstractItemView, QMenu, QDialog, QComboBox, QTextBrowser,
                             QSplitter, QSpinBox, QFrame, QSizePolicy)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QSize, QMimeData, QTimer, QPropertyAnimation
from PyQt6.QtGui import QIcon, QFont, QPalette, QColor, QDragEnterEvent, QDropEvent, QAction

from docx import Document
import openpyxl
import markdown

class ReplacementWorker(QThread):
    progress = pyqtSignal(int)
    file_processed = pyqtSignal(str, bool, int)
    finished = pyqtSignal(dict)

    def __init__(self, files, rules, backup_dir, max_workers=None):
        super().__init__()
        self.files = files
        self.rules = rules
        self.backup_dir = backup_dir
        self.max_workers = max_workers or os.cpu_count()

    def run(self):
        total_files = len(self.files)
        stats = {
            "total_files": total_files,
            "changed_files": 0,
            "total_replacements": 0
        }

        with concurrent.futures.ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            future_to_file = {executor.submit(self.process_file, file_path): file_path for file_path in self.files}
            for i, future in enumerate(concurrent.futures.as_completed(future_to_file)):
                file_path = future_to_file[future]
                try:
                    changed, file_replacements = future.result()
                    if changed:
                        stats["changed_files"] += 1
                        stats["total_replacements"] += file_replacements
                        self.file_processed.emit(file_path, True, file_replacements)
                    else:
                        self.file_processed.emit(file_path, False, 0)
                except Exception as e:
                    self.file_processed.emit(file_path, False, 0)
                self.progress.emit(int((i + 1) / total_files * 100))

        self.finished.emit(stats)

    def process_file(self, file_path):
        backup_path = os.path.join(self.backup_dir, os.path.basename(file_path))
        shutil.copy2(file_path, backup_path)

        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension == '.docx':
            return self.process_word(file_path)
        elif file_extension == '.xlsx':
            return self.process_excel(file_path)
        elif file_extension in ['.txt', '.md']:
            return self.process_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    def process_word(self, file_path):
        doc = Document(file_path)
        changed = False
        file_replacements = 0
        for old_text, new_text in self.rules:
            replacements = self.replace_text_in_document(doc, old_text, new_text)
            if replacements > 0:
                changed = True
                file_replacements += replacements
        if changed:
            doc.save(file_path)
        return changed, file_replacements

    def process_excel(self, file_path):
        changed = False
        file_replacements = 0

        wb = openpyxl.load_workbook(file_path)

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.data_type == 's':
                        original_value = cell.value
                        new_value = original_value
                        for old_text, new_text in self.rules:
                            if old_text in new_value:
                                new_value = new_value.replace(old_text, new_text)
                        if new_value != original_value:
                            cell.value = new_value
                            changed = True
                            file_replacements += 1

        if changed:
            wb.save(file_path)

        return changed, file_replacements

    def process_text(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()

        changed = False
        file_replacements = 0
        for old_text, new_text in self.rules:
            if old_text in content:
                content = content.replace(old_text, new_text)
                changed = True
                file_replacements += content.count(new_text)

        if changed:
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(content)

        return changed, file_replacements

    def replace_text_in_document(self, doc, old_text, new_text):
        replacements = 0
        for para in doc.paragraphs:
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)
                replacements += para.text.count(new_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_text)
                        replacements += cell.text.count(new_text)
        return replacements

class LoadingDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("处理中")
        self.setFixedSize(300, 100)
        layout = QVBoxLayout(self)
        self.label = QLabel("正在处理，请稍候...", self)
        layout.addWidget(self.label)
        self.progress_bar = QProgressBar(self)
        layout.addWidget(self.progress_bar)
        self.setWindowModality(Qt.WindowModality.ApplicationModal)

    def update_progress(self, value):
        self.progress_bar.setValue(value)

class MultiFormatReplacerApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()
        self.replacement_history = []
        self.temp_dir = tempfile.mkdtemp()
        self.file_set = set()

        self.setAcceptDrops(True)

        self.search_timer = QTimer()
        self.search_timer.setSingleShot(True)
        self.search_timer.timeout.connect(self.update_file_list)

        self.is_dark_mode = False
        self.set_style()

    def initUI(self):
        self.setStyle(QStyleFactory.create('Fusion'))
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)

        # 创建顶部工具栏
        toolbar = QToolBar()
        self.addToolBar(Qt.ToolBarArea.TopToolBarArea, toolbar)

        # 添加模式切换按钮到工具栏
        self.mode_switch_button = QPushButton()
        self.mode_switch_button.setIcon(QIcon.fromTheme("weather-clear"))
        self.mode_switch_button.setFixedSize(32, 32)
        self.mode_switch_button.clicked.connect(self.toggle_mode)
        self.mode_switch_button.setStyleSheet("""
            QPushButton {
                border: none;
                border-radius: 16px;
                background-color: transparent;
            }
            QPushButton:hover {
                background-color: rgba(255, 255, 255, 0.1);
            }
        """)
        toolbar.addWidget(self.mode_switch_button)

        toolbar.addSeparator()

        # 在工具栏中添加其他常用操作按钮
        add_file_action = QAction(QIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_FileIcon)), "添加文件", self)
        add_file_action.triggered.connect(self.add_files)
        toolbar.addAction(add_file_action)

        add_folder_action = QAction(QIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_DirIcon)), "添加文件夹", self)
        add_folder_action.triggered.connect(self.add_folder)
        toolbar.addAction(add_folder_action)

        replace_action = QAction(QIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_BrowserReload)), "执行替换", self)
        replace_action.triggered.connect(self.replace_text)
        toolbar.addAction(replace_action)

        undo_action = QAction(QIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_ArrowBack)), "撤销上次替换", self)
        undo_action.triggered.connect(self.undo_last_replacement)
        toolbar.addAction(undo_action)

        # 创建主分割器
        main_splitter = QSplitter(Qt.Orientation.Horizontal)
        main_layout.addWidget(main_splitter)

        # 左侧面板：文件列表和文件操作
        left_panel = QFrame()
        left_panel.setFrameShape(QFrame.Shape.StyledPanel)
        left_layout = QVBoxLayout(left_panel)

        file_label = QLabel("文件列表")
        file_label.setStyleSheet("font-size: 16px; font-weight: bold;")
        left_layout.addWidget(file_label)

        # 添加搜索和过滤布局
        search_filter_layout = QHBoxLayout()

        # 搜索框
        search_layout = QVBoxLayout()
        search_layout.addWidget(QLabel("搜索文件:"))
        self.search_box = QLineEdit()
        self.search_box.setPlaceholderText("输入搜索词 (支持 * 和 ? 通配符)")
        self.search_box.textChanged.connect(self.on_search_text_changed)
        search_layout.addWidget(self.search_box)

        # 文件类型过滤
        filter_layout = QVBoxLayout()
        filter_layout.addWidget(QLabel("文件类型过滤:"))
        self.file_type_filter = QComboBox()
        self.file_type_filter.addItems(["所有文件", "Word (.docx)", "Excel (.xlsx)", "文本 (.txt)", "Markdown (.md)"])
        self.file_type_filter.currentIndexChanged.connect(self.update_file_list)
        filter_layout.addWidget(self.file_type_filter)

        # 将搜索和过滤添加到水平布局
        search_filter_layout.addLayout(search_layout, 2)
        search_filter_layout.addLayout(filter_layout, 1)

        left_layout.addLayout(search_filter_layout)

        self.file_list = QListWidget()
        self.file_list.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        self.file_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.file_list.customContextMenuRequested.connect(self.show_file_list_context_menu)
        left_layout.addWidget(self.file_list)

        file_buttons_layout = QHBoxLayout()
        add_file_button = QPushButton('添加文件')
        add_file_button.setIcon(QIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_FileIcon)))
        add_file_button.clicked.connect(self.add_files)

        add_folder_button = QPushButton('添加文件夹')
        add_folder_button.setIcon(QIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_DirIcon)))
        add_folder_button.clicked.connect(self.add_folder)

        remove_button = QPushButton('移除选中')
        remove_button.setIcon(QIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_TrashIcon)))
        remove_button.clicked.connect(self.remove_selected)

        file_buttons_layout.addWidget(add_file_button)
        file_buttons_layout.addWidget(add_folder_button)
        file_buttons_layout.addWidget(remove_button)
        left_layout.addLayout(file_buttons_layout)

        main_splitter.addWidget(left_panel)
        # 右侧面板：规则、预览和日志
        right_panel = QFrame()
        right_panel.setFrameShape(QFrame.Shape.StyledPanel)
        right_layout = QVBoxLayout(right_panel)
        # 创建右侧垂直分割器
        right_splitter = QSplitter(Qt.Orientation.Vertical)
        right_layout.addWidget(right_splitter)

        # 规则部分
        rules_widget = QWidget()
        rules_layout = QVBoxLayout(rules_widget)
        rules_label = QLabel("替换规则")
        rules_label.setStyleSheet("font-size: 16px; font-weight: bold;")
        rules_layout.addWidget(rules_label)

        self.rules_table = QTableWidget(0, 2)
        self.rules_table.setHorizontalHeaderLabels(['要替换的文本', '新文本'])
        self.rules_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.rules_table.verticalHeader().setDefaultSectionSize(40)  # 增加行高
        self.rules_table.setAlternatingRowColors(True)
        rules_layout.addWidget(self.rules_table)

        rules_buttons_layout = QHBoxLayout()
        add_rule_button = QPushButton('添加规则')
        add_rule_button.setIcon(QIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_FileDialogNewFolder)))
        add_rule_button.clicked.connect(self.add_rule)

        remove_rule_button = QPushButton('删除选中规则')
        remove_rule_button.setIcon(QIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_DialogDiscardButton)))
        remove_rule_button.clicked.connect(self.remove_rule)

        import_rules_button = QPushButton('导入规则')
        import_rules_button.clicked.connect(self.import_rules)
        export_rules_button = QPushButton('导出规则')
        export_rules_button.clicked.connect(self.export_rules)

        rules_buttons_layout.addWidget(add_rule_button)
        rules_buttons_layout.addWidget(remove_rule_button)
        rules_buttons_layout.addWidget(import_rules_button)
        rules_buttons_layout.addWidget(export_rules_button)
        rules_layout.addLayout(rules_buttons_layout)

        right_splitter.addWidget(rules_widget)
        # 并发控制选项
        concurrency_layout = QHBoxLayout()
        concurrency_layout.addWidget(QLabel("最大并发处理数:"))
        self.max_workers_spinbox = QSpinBox()
        self.max_workers_spinbox.setRange(1, os.cpu_count())
        self.max_workers_spinbox.setValue(os.cpu_count())
        concurrency_layout.addWidget(self.max_workers_spinbox)
        rules_layout.addLayout(concurrency_layout)

        # 替换按钮
        replace_button = QPushButton('执行替换')
        replace_button.setIcon(QIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_BrowserReload)))
        replace_button.clicked.connect(self.replace_text)
        rules_layout.addWidget(replace_button)

        # 进度条
        self.progress_bar = QProgressBar()
        rules_layout.addWidget(self.progress_bar)

        # 创建一个分割器来容纳文件预览和操作日志
        preview_log_splitter = QSplitter(Qt.Orientation.Vertical)
        right_splitter.addWidget(preview_log_splitter)

        # 文件预览
        preview_widget = QWidget()
        preview_layout = QVBoxLayout(preview_widget)
        preview_label = QLabel('文件预览')
        preview_label.setStyleSheet("font-size: 16px; font-weight: bold;")
        preview_layout.addWidget(preview_label)
        self.preview_area = QTextBrowser()
        preview_layout.addWidget(self.preview_area)
        preview_log_splitter.addWidget(preview_widget)

        # 操作日志区域
        log_widget = QWidget()
        log_layout = QVBoxLayout(log_widget)
        log_label = QLabel('操作日志')
        log_label.setStyleSheet("font-size: 16px; font-weight: bold;")
        log_layout.addWidget(log_label)
        self.log_area = QTextEdit()
        self.log_area.setReadOnly(True)
        log_layout.addWidget(self.log_area)

        # 添加清理日志按钮
        clear_log_button = QPushButton('清理日志')
        clear_log_button.setIcon(QIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_DialogResetButton)))
        clear_log_button.clicked.connect(self.clear_log)
        log_layout.addWidget(clear_log_button)
        preview_log_splitter.addWidget(log_widget)

        # 撤销按钮
        undo_button = QPushButton('撤销上次替换')
        undo_button.setIcon(QIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_ArrowBack)))
        undo_button.clicked.connect(self.undo_last_replacement)
        right_layout.addWidget(undo_button)

        main_splitter.addWidget(right_panel)

        # 设置初始分割比例
        main_splitter.setSizes([int(self.width() * 0.3), int(self.width() * 0.7)])
        right_splitter.setSizes([int(self.height() * 0.6), int(self.height() * 0.4)])
        preview_log_splitter.setSizes([int(self.height() * 0.5), int(self.height() * 0.5)])

        self.setWindowTitle('多格式文本替换器')
        self.setGeometry(100, 100, 1200, 800)

        # Connect file selection to preview
        self.file_list.itemSelectionChanged.connect(self.update_preview)

    def set_style(self):
        if self.is_dark_mode:
            self.setStyleSheet("""
                QMainWindow, QWidget {
                    background-color: #2c2c2e;
                    color: #ffffff;
                }
                QPushButton {
                    background-color: #3a3a3c;
                    color: #ffffff;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 8px;
                    font-size: 14px;
                }
                QPushButton:hover {
                    background-color: #5a5a5c;
                }
                QPushButton:pressed {
                    background-color: #2a2a2c;
                }
                QTableWidget {
                    background-color: #3a3a3c;
                    color: #ffffff;
                    border: none;
                    gridline-color: #5a5a5c;
                    alternate-background-color: #454547;
                }
                QTableWidget::item:selected {
                    background-color: #007aff;
                }
                QHeaderView::section {
                    background-color: #5a5a5c;
                    color: #ffffff;
                    padding: 8px;
                    border: none;
                }
                QLabel {
                    color: #ffffff;
                    font-size: 14px;
                }
                QComboBox, QLineEdit, QSpinBox {
                    background-color: #3a3a3c;
                    color: #ffffff;
                    border: 1px solid #5a5a5c;
                    padding: 6px;
                    border-radius: 6px;
                }
                QTextEdit, QTextBrowser, QListWidget {
                    background-color: #3a3a3c;
                    color: #ffffff;
                    border: none;
                    border-radius: 8px;
                    padding: 8px;
                }
                QProgressBar {
                    border: none;
                    background-color: #5a5a5c;
                    text-align: center;
                    color: #ffffff;
                }
                QProgressBar::chunk {
                    background-color: #007aff;
                    border-radius: 5px;
                }
            """)
            self.mode_switch_button.setIcon(QIcon.fromTheme("weather-clear-night"))
        else:
            self.setStyleSheet("""
                QMainWindow, QWidget {
                    background-color: #f0f0f0;
                    color: #000000;
                }
                QPushButton {
                    background-color: #e0e0e0;
                    color: #000000;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 8px;
                    font-size: 14px;
                }
                QPushButton:hover {
                    background-color: #d0d0d0;
                }
                QPushButton:pressed {
                    background-color: #c0c0c0;
                }
                QTableWidget {
                    background-color: #ffffff;
                    color: #000000;
                    border: 1px solid #d0d0d0;
                    gridline-color: #e0e0e0;
                    alternate-background-color: #f5f5f5;
                }
                QTableWidget::item:selected {
                    background-color: #3498db;
                    color: #ffffff;
                }
                QHeaderView::section {
                    background-color: #e0e0e0;
                    color: #000000;
                    padding: 8px;
                    border: none;
                }
                QLabel {
                    color: #000000;
                    font-size: 14px;
                }
                QComboBox, QLineEdit, QSpinBox {
                    background-color: #ffffff;
                    color: #000000;
                    border: 1px solid #d0d0d0;
                    padding: 6px;
                    border-radius: 6px;
                }
                QTextEdit, QTextBrowser, QListWidget {
                    background-color: #ffffff;
                    color: #000000;
                    border: 1px solid #d0d0d0;
                    border-radius: 8px;
                    padding: 8px;
                }
                QProgressBar {
                    border: none;
                    background-color: #e0e0e0;
                    text-align: center;
                    color: #000000;
                }
                QProgressBar::chunk {
                    background-color: #3498db;
                    border-radius: 5px;
                }
            """)
            self.mode_switch_button.setIcon(QIcon.fromTheme("weather-clear"))
    def toggle_mode(self):
        self.is_dark_mode = not self.is_dark_mode
        self.set_style()

    def show_file_list_context_menu(self, position):
        menu = QMenu()
        menu.setStyleSheet("""
            QMenu {
                background-color: #3a3a3c;
                color: #ffffff;
                border: none;
            }
            QMenu::item:selected {
                background-color: #5a5a5c;
            }
        """ if self.is_dark_mode else """
            QMenu {
                background-color: #ffffff;
                color: #000000;
                border: 1px solid #d0d0d0;
            }
            QMenu::item:selected {
                background-color: #e0e0e0;
            }
        """)
        remove_action = menu.addAction("移除选中")
        open_folder_action = menu.addAction("在资源管理器中打开")

        action = menu.exec(self.file_list.mapToGlobal(position))
        if action == remove_action:
            self.remove_selected()
        elif action == open_folder_action:
            self.open_selected_in_explorer()

    def open_selected_in_explorer(self):
        selected_items = self.file_list.selectedItems()
        if not selected_items:
            return
        file_path = selected_items[0].text()
        os.startfile(os.path.dirname(file_path))

    def add_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "选择文件", "", "所有支持的文件 (*.docx *.xlsx *.txt *.md)")
        new_files = [file for file in files if file not in self.file_set]
        if new_files:
            self.file_list.addItems(new_files)
            self.file_set.update(new_files)
            self.log(f"已添加 {len(new_files)} 个新文件。")
        if len(new_files) < len(files):
            self.log(f"已跳过 {len(files) - len(new_files)} 个重复文件。")
        self.update_file_list()

    def add_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "选择文件夹")
        if folder:
            new_files = []
            for root, dirs, files in os.walk(folder):
                for file in files:
                    if file.endswith(('.docx', '.xlsx', '.txt', '.md')):
                        file_path = os.path.join(root, file)
                        if file_path not in self.file_set:
                            new_files.append(file_path)
                            self.file_set.add(file_path)
            self.file_list.addItems(new_files)
            self.log(f"已从文件夹添加 {len(new_files)} 个新文件。")
            if len(new_files) < len([f for f in os.listdir(folder) if f.endswith(('.docx', '.xlsx', '.txt', '.md'))]):
                self.log("部分文件因重复而被跳过。")
            self.update_file_list()

    def remove_selected(self):
        selected_items = self.file_list.selectedItems()
        if not selected_items:
            return

        # 暂时阻塞信号以避免频繁更新
        self.file_list.blockSignals(True)

        # 创建要移除的项目的行号列表
        rows_to_remove = sorted([self.file_list.row(item) for item in selected_items], reverse=True)

        # 批量移除项目
        for row in rows_to_remove:
            item = self.file_list.takeItem(row)
            self.file_set.remove(item.text())

        # 恢复信号
        self.file_list.blockSignals(False)

        # 更新一次UI
        self.file_list.update()

        self.log(f"已移除 {len(selected_items)} 个文件。")

        # 如果文件列表变空，更新文件类型过滤器
        if self.file_list.count() == 0:
            self.update_file_list()

    def add_rule(self, old_text="", new_text=""):
        row_position = self.rules_table.rowCount()
        self.rules_table.insertRow(row_position)
        self.rules_table.setItem(row_position, 0, QTableWidgetItem(old_text))
        self.rules_table.setItem(row_position, 1, QTableWidgetItem(new_text))
        self.validate_rules()

    def remove_rule(self):
        indices = self.rules_table.selectionModel().selectedRows()
        for index in sorted(indices, reverse=True):
            self.rules_table.removeRow(index.row())
        self.validate_rules()

    def validate_rules(self):
        rules = set()
        for row in range(self.rules_table.rowCount()):
            old_text = self.rules_table.item(row, 0).text().strip()
            new_text = self.rules_table.item(row, 1).text().strip()
            if old_text and new_text:
                if old_text == new_text:
                    self.rules_table.item(row, 0).setBackground(QColor(255, 100, 100))
                    self.rules_table.item(row, 1).setBackground(QColor(255, 100, 100))
                elif (old_text, new_text) in rules:
                    self.rules_table.item(row, 0).setBackground(QColor(255, 200, 100))
                    self.rules_table.item(row, 1).setBackground(QColor(255, 200, 100))
                else:
                    self.rules_table.item(row, 0).setBackground(QColor(60, 60, 62) if self.is_dark_mode else QColor(255, 255, 255))
                    self.rules_table.item(row, 1).setBackground(QColor(60, 60, 62) if self.is_dark_mode else QColor(255, 255, 255))
                    rules.add((old_text, new_text))
            else:
                self.rules_table.item(row, 0).setBackground(QColor(60, 60, 62) if self.is_dark_mode else QColor(255, 255, 255))
                self.rules_table.item(row, 1).setBackground(QColor(60, 60, 62) if self.is_dark_mode else QColor(255, 255, 255))

    def replace_text(self):
        files = [self.file_list.item(i).text() for i in range(self.file_list.count())]
        rules = []
        for row in range(self.rules_table.rowCount()):
            old_text = self.rules_table.item(row, 0).text().strip()
            new_text = self.rules_table.item(row, 1).text().strip()
            if old_text and new_text and old_text != new_text:
                rules.append((old_text, new_text))

        if not files or not rules:
            self.show_styled_message_box("警告", "请添加文件和有效的替换规则。", QMessageBox.Icon.Warning)
            return

        confirm = self.show_styled_message_box('确认', f'是否执行替换操作？\n文件数：{len(files)}\n规则数：{len(rules)}',
                                               QMessageBox.Icon.Question,
                                               QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if confirm == QMessageBox.StandardButton.No:
            return

        backup_dir = os.path.join(self.temp_dir, f'backup_{len(self.replacement_history)}')
        os.makedirs(backup_dir, exist_ok=True)

        self.log(f"开始替换操作：处理 {len(files)} 个文件，应用 {len(rules)} 条规则。")

        max_workers = self.max_workers_spinbox.value()
        self.worker = ReplacementWorker(files, rules, backup_dir, max_workers)
        self.worker.progress.connect(self.update_progress)
        self.worker.file_processed.connect(self.update_output)
        self.worker.finished.connect(self.replacement_finished)

        self.loading_dialog = LoadingDialog(self)
        self.worker.progress.connect(self.loading_dialog.update_progress)
        self.loading_dialog.show()

        self.worker.start()

        self.replacement_history.append((files, rules, backup_dir))
    def update_progress(self, value):
        if self.progress_bar.value() < value:
            animation = QPropertyAnimation(self.progress_bar, b"value")
            animation.setDuration(200)  # 200毫秒的动画时长
            animation.setStartValue(self.progress_bar.value())
            animation.setEndValue(value)
            animation.start()
        else:
            self.progress_bar.setValue(value)

        if value == 100:
            self.progress_bar.setStyleSheet("""
                QProgressBar {
                    border: none;
                    background-color: #5a5a5c;
                    text-align: center;
                    color: #ffffff;
                }
                QProgressBar::chunk {
                    background-color: #4cd964;
                    border-radius: 5px;
                }
            """)
        else:
            self.progress_bar.setStyleSheet("""
                QProgressBar {
                    border: none;
                    background-color: #5a5a5c;
                    text-align: center;
                    color: #ffffff;
                }
                QProgressBar::chunk {
                    background-color: #007aff;
                    border-radius: 5px;
                }
            """)

    def update_output(self, file_path, changed, replacements):
        status = f"已替换 ({replacements} 处)" if changed else "未更改"
        self.log(f"{file_path}: {status}")

    def replacement_finished(self, stats):
        self.loading_dialog.close()
        self.log("替换操作完成。")
        self.progress_bar.setValue(100)

        summary = (f"替换操作摘要:\n"
                   f"处理文件总数: {stats['total_files']}\n"
                   f"发生更改的文件数: {stats['changed_files']}\n"
                   f"总替换次数: {stats['total_replacements']}")

        self.show_styled_message_box("替换完成", summary, QMessageBox.Icon.Information)

    def undo_last_replacement(self):
        if not self.replacement_history:
            self.show_styled_message_box("提示", "没有可撤销的操作。", QMessageBox.Icon.Information)
            return

        files, rules, backup_dir = self.replacement_history.pop()
        confirm = self.show_styled_message_box('确认', f'是否撤销上次替换操作？\n文件数：{len(files)}\n规则数：{len(rules)}',
                                               QMessageBox.Icon.Question,
                                               QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if confirm == QMessageBox.StandardButton.No:
            return

        self.log("开始撤销上次替换操作...")
        for file_path in files:
            backup_path = os.path.join(backup_dir, os.path.basename(file_path))
            if os.path.exists(backup_path):
                try:
                    shutil.copy2(backup_path, file_path)
                    self.log(f"已恢复文件: {file_path}")
                except Exception as e:
                    self.log(f"恢复文件失败: {file_path}, 错误: {str(e)}")
            else:
                self.log(f"未找到备份文件: {file_path}")

        self.log("撤销操作完成。")
        self.show_styled_message_box("撤销完成", "已成功撤销上次替换操作。", QMessageBox.Icon.Information)

    def update_preview(self):
        selected_items = self.file_list.selectedItems()
        if not selected_items:
            self.preview_area.clear()
            return

        file_path = selected_items[0].text()
        file_extension = os.path.splitext(file_path)[1].lower()

        try:
            if file_extension == '.docx':
                doc = Document(file_path)
                preview_text = '\n'.join([para.text for para in doc.paragraphs[:20]])
            elif file_extension == '.xlsx':
                wb = openpyxl.load_workbook(file_path, read_only=True)
                sheet = wb.active
                preview_text = '\n'.join([', '.join(str(cell.value) for cell in row) for row in sheet.iter_rows(max_row=20)])
            elif file_extension in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as file:
                    preview_text = file.read(2000)  # Read first 2000 characters
            else:
                preview_text = "不支持的文件类型"

            self.preview_area.setText(preview_text)
        except Exception as e:
            self.preview_area.setText(f"无法预览文件: {str(e)}")

    def import_rules(self):
        file_name, _ = QFileDialog.getOpenFileName(self, "导入规则", "", "JSON Files (*.json)")
        if file_name:
            try:
                with open(file_name, 'r', encoding='utf-8') as f:
                    imported_rules = json.load(f)

                # 获取现有规则
                existing_rules = []
                for row in range(self.rules_table.rowCount()):
                    old_text = self.rules_table.item(row, 0).text().strip()
                    new_text = self.rules_table.item(row, 1).text().strip()
                    if old_text and new_text:
                        existing_rules.append((old_text, new_text))

                # 合并规则，保留现有规则，更新重复的规则
                merged_rules = dict(existing_rules)
                for old_text, new_text in imported_rules:
                    merged_rules[old_text] = new_text

                # 清空规则表格并添加合并后的规则
                self.rules_table.setRowCount(0)
                for old_text, new_text in merged_rules.items():
                    self.add_rule(old_text, new_text)

                self.log(f"已从 {file_name} 导入并合并 {len(imported_rules)} 条规则")
            except Exception as e:
                self.show_styled_message_box("导入失败", f"导入规则失败: {str(e)}", QMessageBox.Icon.Warning)

    def export_rules(self):
        rules = []
        for row in range(self.rules_table.rowCount()):
            old_text = self.rules_table.item(row, 0).text().strip()
            new_text = self.rules_table.item(row, 1).text().strip()
            if old_text and new_text:
                rules.append([old_text, new_text])

        if not rules:
            self.show_styled_message_box("导出失败", "没有规则可以导出", QMessageBox.Icon.Warning)
            return

        file_name, _ = QFileDialog.getSaveFileName(self, "导出规则", "", "JSON Files (*.json)")
        if file_name:
            try:
                with open(file_name, 'w', encoding='utf-8') as f:
                    json.dump(rules, f, ensure_ascii=False, indent=2)
                self.log(f"已将 {len(rules)} 条规则导出到 {file_name}")
            except Exception as e:
                self.show_styled_message_box("导出失败", f"导出规则失败: {str(e)}", QMessageBox.Icon.Warning)

    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event: QDropEvent):
        files = [u.toLocalFile() for u in event.mimeData().urls()]
        self.process_dropped_files(files)

    def process_dropped_files(self, files):
        new_files = []
        for file in files:
            if os.path.isdir(file):
                for root, dirs, files in os.walk(file):
                    for f in files:
                        if f.endswith(('.docx', '.xlsx', '.txt', '.md')):
                            file_path = os.path.join(root, f)
                            if file_path not in self.file_set:
                                new_files.append(file_path)
                                self.file_set.add(file_path)
            elif file.endswith(('.docx', '.xlsx', '.txt', '.md')) and file not in self.file_set:
                new_files.append(file)
                self.file_set.add(file)

        self.file_list.addItems(new_files)
        self.log(f"通过拖放添加了 {len(new_files)} 个新文件。")
        if len(new_files) < len(files):
            self.log(f"跳过了 {len(files) - len(new_files)} 个重复或不支持的文件。")
        self.update_file_list()

    def log(self, message):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.log_area.append(f"[{timestamp}] {message}")

    def clear_log(self):
        self.log_area.clear()
        self.log("日志已清理。")

    def update_file_list(self):
        filter_text = self.file_type_filter.currentText()
        search_text = self.search_box.text().lower()
        self.file_list.clear()

        for file in self.file_set:
            file_name = os.path.basename(file).lower()
            if (filter_text == "所有文件" or file.endswith(filter_text.split('.')[-1].strip(')'))) and \
                    (not search_text or fnmatch.fnmatch(file_name, f"*{search_text}*")):
                self.file_list.addItem(file)

    def on_search_text_changed(self):
        # 当搜索文本改变时，启动计时器
        self.search_timer.start(300)  # 300毫秒后触发搜索

    def show_styled_message_box(self, title, text, icon=QMessageBox.Icon.Information, buttons=QMessageBox.StandardButton.Ok):
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle(title)
        msg_box.setText(text)
        msg_box.setIcon(icon)
        msg_box.setStandardButtons(buttons)

        if self.is_dark_mode:
            msg_box.setStyleSheet("""
                QMessageBox {
                    background-color: #2c2c2e;
                }
                QMessageBox QLabel {
                    color: #ffffff;
                }
                QPushButton {
                    background-color: #3a3a3c;
                    color: #ffffff;
                    border: none;
                    padding: 5px 15px;
                    border-radius: 3px;
                }
                QPushButton:hover {
                    background-color: #5a5a5c;
                }
            """)
        else:
            msg_box.setStyleSheet("""
                QMessageBox {
                    background-color: #f0f0f0;
                }
                QMessageBox QLabel {
                    color: #000000;
                }
                QPushButton {
                    background-color: #e0e0e0;
                    color: #000000;
                    border: none;
                    padding: 5px 15px;
                    border-radius: 3px;
                }
                QPushButton:hover {
                    background-color: #d0d0d0;
                }
            """)

        return msg_box.exec()

    def closeEvent(self, event):
        # 关闭应用时清理临时目录
        shutil.rmtree(self.temp_dir, ignore_errors=True)
        super().closeEvent(event)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setStyle(QStyleFactory.create('Fusion'))
    ex = MultiFormatReplacerApp()
    ex.show()
    sys.exit(app.exec())